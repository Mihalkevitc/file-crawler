import os
import random
import zipfile
import rarfile  # для .rar
import py7zr    # для .7z
from faker import Faker
from docx import Document
from openpyxl import Workbook
from fpdf import FPDF  # для .pdf
import patoolib
import subprocess

fake = Faker('ru_RU')

def create_data_directory():
    """Создаём папки для хранения тестовых файлов"""
    os.makedirs('data/docs', exist_ok=True)
    os.makedirs('data/archives', exist_ok=True)

def generate_text():
    """Генерируем случайный текст для наполнения документов"""
    topics = ['финансы', 'инвестиции', 'налоги', 'бухгалтерия', 'аудит']
    topic = random.choice(topics)
    return f"""
{fake.catch_phrase()}
Тема: {topic}
Дата: {fake.date()}
Автор: {fake.name()}
Компания: {fake.company()}
{fake.text(max_nb_chars=300)}
"""

def create_doc(path):
    """Создаём .doc файл (просто текст, так как .doc сложный формат)"""
    with open(path, 'w', encoding='utf-8') as f:
        f.write(generate_text())

def create_docx(path):
    """Создаём .docx файл"""
    doc = Document()
    doc.add_heading(fake.catch_phrase(), level=1)
    doc.add_paragraph(generate_text())
    doc.save(path)

def create_xls(path):
    """Создаём .xls файл (просто текст)"""
    with open(path, 'w', encoding='utf-8') as f:
        f.write(generate_text())

def create_xlsx(path):
    """Создаём .xlsx файл"""
    wb = Workbook()
    ws = wb.active
    ws.append(['Дата', 'Документ', 'Сумма', 'Контрагент'])
    for _ in range(random.randint(3, 7)):
        ws.append([fake.date(), fake.catch_phrase(), 
                   random.randint(1000, 500000), fake.company()])
    wb.save(path)

def create_pdf(path):
    """Создаём .pdf файл (упрощённо)"""
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font('Arial', size=12)
    pdf.multi_cell(200, 10, txt=generate_text())
    pdf.output(path)

def create_txt(path):
    """Создаём .txt файл"""
    with open(path, 'w', encoding='utf-8') as f:
        f.write(generate_text())

def create_archive_zip():
    """Создаём .zip архив"""
    name = f"data/archives/archive_{fake.word()}.zip"
    with zipfile.ZipFile(name, 'w') as zf:
        for i in range(3):
            temp = f"temp_{i}.txt"
            create_txt(temp)
            zf.write(temp)
            os.remove(temp)
    print(f"  Создан архив: {name}")

def create_archive_rar():
    """Создаём .rar архив с помощью patool (требуется WinRAR)"""
    name = f"data/archives/archive_{fake.word()}.rar"
    
    # Создание временных файлов
    temp_files = []
    for i in range(3):
        temp = f"temp_{i}.txt"
        create_txt(temp)
        temp_files.append(temp)
    
    try:
        # Пытаемся создать RAR
        patoolib.create_archive(name, temp_files)
        print(f"  Создан архив: {name}")
    except Exception as e:
        print(f"  Не удалось создать RAR: {e}")
        print(f"  Пропускаем RAR архив")
    finally:
        # Удаляем временные файлы
        for temp in temp_files:
            if os.path.exists(temp):
                os.remove(temp)

def create_archive_7z():
    """Создаём .7z архив"""
    name = f"data/archives/archive_{fake.word()}.7z"
    with py7zr.SevenZipFile(name, 'w') as sz:
        for i in range(3):
            temp = f"temp_{i}.txt"
            create_txt(temp)
            sz.write(temp)
            os.remove(temp)
    print(f"  Создан архив: {name}")

def main():
    create_data_directory()
    for i in range(3):  # по 3 файла каждого типа
        create_doc(f"data/docs/doc_{i}.doc")
        create_docx(f"data/docs/doc_{i}.docx")
        create_xls(f"data/docs/doc_{i}.xls")
        create_xlsx(f"data/docs/doc_{i}.xlsx")
        # create_pdf(f"data/docs/doc_{i}.pdf")
        create_txt(f"data/docs/doc_{i}.txt")
    
    # Создаём по одному архиву каждого типа
    create_archive_zip()
    create_archive_rar()
    create_archive_7z()

if __name__ == "__main__":
    main()